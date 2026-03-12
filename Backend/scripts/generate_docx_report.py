import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Configuration du titre principal
    title = doc.add_heading('RAPPORT TECHNIQUE DÉTAILLÉ : HOTELY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Date : Mars 2026', style='Subtitle')
    doc.add_paragraph('Projet : Plateforme de Réservation Hotely', style='Subtitle')
    
    doc.add_page_break()

    # 1. Résumé Exécutif
    doc.add_heading('1. Résumé Exécutif', level=1)
    doc.add_paragraph(
        "Hotely est une application web full-stack innovante conçue "
        "pour simplifier la réservation hôtelière. Elle connecte les voyageurs avec "
        "les propriétaires d'hôtels via une interface fluide et sécurisée."
    )
    doc.add_paragraph(
        "Ce rapport détaille l'architecture globale, la base de données, l'API backend, "
        "l'organisation du code frontend, ainsi que les fonctionnalités avancées comme le "
        "chatbot IA et le système de notifications en temps réel."
    )

    # 2. Fonctionnalités Principales
    doc.add_heading('2. Fonctionnalités Principales', level=1)
    features = [
        "Authentification sécurisée (Email/Mot de passe, Google Sign-In) via Firebase et JWT.",
        "Vérification par code OTP (Email) et réinitialisation sécurisée des mots de passe.",
        "Recherche intuitive d'hôtels et de chambres avec détails complets et galeries d'images.",
        "Réservation de chambre avec calcul automatique du prix total selon les dates.",
        "Passerelles de paiement locales intégrées (Bankily, Sedad, Masrivi, ou Carte Bancaire) avec preuve d'achat.",
        "Tableau de bord propriétaire robuste incluant statistiques, revenus, vue calendrier et analytiques.",
        "Espace d'administration centralisé (Control Center) pour modérer l'ensemble de la plateforme.",
        "Messagerie privée et notifications temps réel.",
        "Chatbot d'assistance propulsé par l'IA (RAG - LLaMA 3.1) capable de répondre aux questions des utilisateurs."
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')

    # 3. Stack Technologique
    doc.add_heading('3. Stack Technologique', level=1)
    
    doc.add_heading('Backend', level=2)
    backend_techs = [
        ("Python 3 & Flask", "Framework API et logique métier."),
        ("SQLAlchemy & Alembic", "ORM et gestion des migrations de base de données."),
        ("JWT Extended", "Gestion des tokens pour la sécurité des endpoints."),
        ("LangChain & Groq", "Pour le Chatbot IA basé sur un modèle LLaMA 3.1."),
        ("ChromaDB & FastEmbed", "Base vectorielle et embeddings pour le RAG."),
        ("SQLite / MySQL", "Base de données relationnelle.")
    ]
    for tech, desc in backend_techs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tech + " : ").bold = True
        p.add_run(desc)

    doc.add_heading('Frontend', level=2)
    frontend_techs = [
        ("React 18 & Vite", "Framework UI pour une Single Page Application réactive."),
        ("React Router DOM", "Navigation côté client sans rechargement de page."),
        ("Firebase Auth", "Outil de vérification et gestion des fournisseurs d'identité."),
        ("Axios", "Client HTTP pour interroger l'API Flask.")
    ]
    for tech, desc in frontend_techs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tech + " : ").bold = True
        p.add_run(desc)

    # 4. Modèles de Base de Données
    doc.add_heading('4. Modèles de Données Principaux', level=1)
    
    models_info = [
        ("User", "Stocke les clients, managers et administrateurs. Contient les hash, tokens OTP et permissions de rôle."),
        ("Hotel & Room", "Définissent les établissements et les chambres. Gèrent la disponibilité, les prix, et les équipements."),
        ("Booking", "Historise les réservations de la date d'arrivée jusqu'à la date de départ avec une référence de commande (ex: BK...)."),
        ("Payment", "Supervise l'état du paiement d'une réservation, avec gestion des méthodes de paiements et captures d'écran."),
        ("Review & Message", "Garantissent l'interaction client/propriétaire avec un système de notation et une boîte de réception.")
    ]
    for name, desc in models_info:
        p = doc.add_paragraph()
        p.add_run(f"Modèle {name}").bold = True
        doc.add_paragraph(desc)

    # 5. Endpoints API
    doc.add_heading('5. Architecture de l\'API Backend', level=1)
    doc.add_paragraph("L'API est organisée par Blueprints au niveau du backend :")
    
    api_routes = [
        ("/api/auth/*", "Gestion des connexions, OTP, profil et admins (Firebase + JWT)."),
        ("/api/bookings/*", "Créer, lister, annuler et confirmer des réservations hôtelières."),
        ("/api/dashboard/*", "Récupérer des agrégations avancées (Revenus, Statistiques, Remplissage)."),
        ("/api/chat", "Endpoint dédié au chatbot IA question-réponse RAG.")
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chemin API'
    hdr_cells[1].text = 'Responsabilité'
    for route, resp in api_routes:
        row_cells = table.add_row().cells
        row_cells[0].text = route
        row_cells[1].text = resp

    # 6. Structuration Frontend
    doc.add_heading('6. Architecture Frontend (React)', level=1)
    doc.add_paragraph("L'interface utilisateur est architecturée pour la lisibilité et l'évolutivité avec 16 pages et "
                      "25 composants :")
    pages = [
        ("Pages Publiques", "Home, HotelDetails, RoomDetails."),
        ("Pages d'Authentification", "Login, Register, ForgotPassword, ResetPassword, EmailVerification."),
        ("Pages Utilisateur", "Profile, Settings, Messages, Notifications, Payment."),
        ("Pages Privilégiées", "Dashboard (Managers), Control Center (Admins).")
    ]
    for cat, pgs in pages:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cat + " - ").bold = True
        p.add_run(pgs)

    # 7. Sécurité et Flux Logistique
    doc.add_heading('7. Sécurité et Contrôle d\'Accès', level=1)
    doc.add_paragraph("Un système de permissions strict a été mis en place :")
    p1 = doc.add_paragraph("1. Séparation des rôles : ", style='List Number')
    p1.add_run("Clients, Managers et Admins.").bold = True
    p1.add_run(" Le propriétaire gère ses propres statistiques et hôtels de manière cloisonnée.")
    
    p2 = doc.add_paragraph("2. Sécurité OTP et Firebase : ", style='List Number')
    p2.add_run("Validation à deux facteurs pour l'enregistrement et récupération sécurisée avec Firebase SDK.")
    
    p3 = doc.add_paragraph("3. JWT Access & Refresh : ", style='List Number')
    p3.add_run("Les endpoints sensibles nécessitent un token actif renouvelable pour limiter l'exposition.")

    # 8. Intelligence Artificielle et Support
    doc.add_heading('8. Intégration Intelligence Artificielle', level=1)
    doc.add_paragraph(
        "L'application inclut un système RAG (Retrieval-Augmented Generation) exclusif."
        " Le composant interactif `FloatingChatbot.jsx` permet de requêter le NLP (LLaMA 3.1) "
        "via l'API Groq."
    )
    doc.add_paragraph(
        "Ce chatbot est contextuellement alimenté (ChromaDB Vector Store) "
        "par les règles de l'établissement à partir d'un document maître."
    )

    doc.save('/home/abass/Desktop/vibepi/Hotely_Rapport_Technique.docx')

if __name__ == '__main__':
    main()
